import os
import json
import re
from docx import Document

VISIT_FILES = [
    r"c:\$$$\TENS11\TEN11-MARCH2026.docx",
    r"c:\$$$\TENS11\TEN11-APRIL2026.docx",
    r"c:\$$$\TENS11\TEN11-MAY 2026.docx",
    r"c:\$$$\TENS11\TEN11-JUNE 2026.docx"
]
PRODUCT_FILE = r"c:\$$$\TENS11\PRODUCT SELL -TEN11 - 2026.docx"

def clean_amount(amount_str):
    match = re.search(r'\d+', amount_str.replace(',', ''))
    return int(match.group()) if match else 0

def clean_payment_method(amount_str):
    if 'UPI' in amount_str.upper():
        return 'UPI'
    if 'CARD' in amount_str.upper():
        return 'Card'
    if 'BANK' in amount_str.upper():
        return 'Bank Transfer'
    return 'Cash'

def clean_phone(phone_str):
    phone = re.sub(r'[^\d]', '', phone_str)
    if len(phone) >= 10:
        return phone[-10:] # get last 10 digits
    return phone

def parse_visits():
    all_visits = []
    
    for filepath in VISIT_FILES:
        try:
            doc = Document(filepath)
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0: continue # Skip header
                    
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    if len(cells) < 9: continue # Bad row
                    
                    if not cells[1]: continue # No date
                    
                    date_str = cells[1]
                    name = cells[2]
                    phone = clean_phone(cells[3])
                    services = cells[6]
                    staff = cells[7]
                    amount_raw = cells[8]
                    
                    if not name or name == '-': continue
                    
                    all_visits.append({
                        "source": os.path.basename(filepath),
                        "date": date_str,
                        "customer_name": name,
                        "customer_phone": phone,
                        "services": services,
                        "staff": staff,
                        "amount": clean_amount(amount_raw),
                        "payment_method": clean_payment_method(amount_raw)
                    })
        except Exception as e:
            print(f"Failed to read {filepath}: {e}")
            
    with open('clean_visits.json', 'w', encoding='utf-8') as f:
        json.dump(all_visits, f, indent=2)
    print(f"Extracted {len(all_visits)} visits.")

def parse_products():
    all_sales = []
    
    try:
        doc = Document(PRODUCT_FILE)
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0: continue # Skip header
                
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                if len(cells) < 6: continue
                
                if not cells[0]: continue
                
                date_str = cells[0]
                name = cells[1]
                phone = clean_phone(cells[2])
                product = cells[3]
                amount_raw = cells[4]
                staff = cells[5]
                
                if name == '-': 
                    name = "Retail Walk-in"
                
                all_sales.append({
                    "date": date_str,
                    "customer_name": name,
                    "customer_phone": phone,
                    "product": product,
                    "staff": staff,
                    "amount": clean_amount(amount_raw),
                    "payment_method": clean_payment_method(amount_raw)
                })
    except Exception as e:
        print(f"Failed to read {PRODUCT_FILE}: {e}")
        
    with open('clean_products.json', 'w', encoding='utf-8') as f:
        json.dump(all_sales, f, indent=2)
    print(f"Extracted {len(all_sales)} product sales.")

if __name__ == "__main__":
    parse_visits()
    parse_products()
